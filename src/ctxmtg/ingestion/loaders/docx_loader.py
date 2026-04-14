# Copyright (c) 2026 Aliud Inquisito Inc. / Dhananjay Raol. All rights reserved.
# Proprietary and confidential. See LICENSE file for details.
"""
DOCX File Loader
================

Loads .docx files using python-docx. Extracts paragraph text and
basic table content. Images and complex formatting are discarded --
only the textual content matters for knowledge extraction.

Depends on:
    - docx (python-docx package)
    - ctxmtg.models.interaction (Interaction, SourceType)
    - ctxmtg.storage.id_gen (generate_interaction_id)
    - ctxmtg.exceptions (IntakeError)

Used by:
    - ctxmtg.ingestion.loaders (registered for .docx extension)
"""

from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path

import structlog

from ctxmtg.exceptions import IntakeError
from ctxmtg.models.interaction import Interaction, SourceType
from ctxmtg.storage.id_gen import generate_interaction_id

logger = structlog.get_logger("ctxmtg.ingestion.loaders.docx_loader")


def load_docx_file(file_path: Path) -> Interaction:
    """
    Load a .docx file and produce an Interaction object.

    Extracts all paragraph text and table cell text. Paragraphs are
    joined with newlines. Table rows are formatted as pipe-separated
    values.

    Args:
        file_path: Path to the .docx file.

    Returns:
        An Interaction containing the document text.

    Raises:
        IntakeError: If python-docx is not installed, the file cannot
            be read, or the document has no text content.
    """
    if not file_path.exists():
        raise IntakeError(f"DOCX file not found: {file_path}", error_code="CTXMTG-ING-001")
    if not file_path.is_file():
        raise IntakeError(f"Path is not a file: {file_path}", error_code="CTXMTG-ING-001")

    try:
        import docx
    except ImportError as exc:
        raise IntakeError(
            "python-docx is required for .docx files: pip install python-docx",
            error_code="CTXMTG-ING-006",
        ) from exc

    try:
        doc = docx.Document(str(file_path))
    except Exception as exc:
        raise IntakeError(
            f"Cannot open DOCX file {file_path}: {exc}", error_code="CTXMTG-ING-001"
        ) from exc

    parts: list[str] = []

    # Extract paragraph text.
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract table content.
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    content = "\n".join(parts)
    if not content.strip():
        raise IntakeError(
            f"DOCX file has no extractable text: {file_path}",
            error_code="CTXMTG-ING-003",
        )

    # Try to get the document title from core properties.
    title = file_path.stem
    try:
        if doc.core_properties.title:
            title = doc.core_properties.title
    except Exception:
        pass

    interaction_id = generate_interaction_id("doc", content)

    logger.info("docx_file_loaded", file_path=str(file_path), paragraphs=len(parts))

    return Interaction(
        id=interaction_id,
        source_type=SourceType.DOC,
        title=title,
        content=content,
        participants=[],
        metadata={"source_file": file_path.name, "source_format": "docx"},
        created_at=datetime.now(timezone.utc),
    )
